from pathlib import Path

import pytest
from docx import Document

from app.cv.parsing_service import order_pdf_columns
from app.cv.parsing_service import CVParsingError
from app.cv.parsing_service import calculate_pdf_extraction_quality
from app.cv.parsing_service import extract_pdf_page_layout_text
from app.cv.parsing_service import extract_section_lines
from app.cv.parsing_service import extract_text_from_cv
from app.cv.parsing_service import extract_text_from_docx
from app.cv.parsing_service import extract_text_from_pdf
from app.cv.parsing_service import extract_text_from_pdf_pdfplumber
from app.cv.parsing_service import extract_text_from_pdf_pypdf2
from app.cv.parsing_service import normalize_text_lines
from app.cv.parsing_service import parse_cv_text
from app.cv.parsing_service import select_best_pdf_extraction
from app.cv.parsing_service import extract_languages
from app.cv.parsing_service import merge_split_section_headings


CV_PARSER_FIXTURES = (
    Path(__file__).parent
    / "fixtures"
    / "cv_parser"
)


def test_parse_cv_text_extracts_basic_profile_information():
    raw_text = """
Vincent Gueret
Technical Partnerships Manager

Summary
Technical partnerships profile with API integration experience.

Skills
Python, FastAPI, React, PostgreSQL

Languages
French, English, Portuguese

Certifications
Azure Fundamentals
Scrum

Experience
Technical Partnerships Manager at Software Company
Managed marketplace, webstore and carrier integrations.
"""

    parsed_cv = parse_cv_text(raw_text)

    assert parsed_cv.full_name == "Vincent Gueret"
    assert parsed_cv.professional_title == "Technical Partnerships Manager"
    assert parsed_cv.summary == (
        "Technical partnerships profile with API integration experience."
    )
    assert "Python" in parsed_cv.skills
    assert "FastAPI" in parsed_cv.skills
    assert "React" in parsed_cv.skills
    assert "PostgreSQL" in parsed_cv.skills
    assert "French" in parsed_cv.languages
    assert "English" in parsed_cv.languages
    assert "Portuguese" in parsed_cv.languages
    assert "Azure Fundamentals" in parsed_cv.certifications
    assert "Scrum" in parsed_cv.certifications
    assert len(parsed_cv.experiences) >= 1


def test_parse_cv_text_returns_empty_lists_when_sections_are_missing():
    raw_text = """
Vincent Gueret
Technical Partnerships Manager
"""

    parsed_cv = parse_cv_text(raw_text)

    assert parsed_cv.full_name == "Vincent Gueret"
    assert parsed_cv.professional_title == "Technical Partnerships Manager"
    assert parsed_cv.summary is None
    assert parsed_cv.skills == []
    assert parsed_cv.languages == []
    assert parsed_cv.certifications == []
    assert parsed_cv.experiences == []


def test_extract_text_from_cv_rejects_unsupported_file_format(
    tmp_path: Path,
):
    unsupported_file = tmp_path / "cv.txt"
    unsupported_file.write_text(
        "Plain text CV",
        encoding="utf-8",
    )

    with pytest.raises(CVParsingError) as exc_info:
        extract_text_from_cv(unsupported_file)

    assert "Unsupported CV file format" in str(exc_info.value)


def test_extract_text_from_cv_rejects_missing_file(
    tmp_path: Path,
):
    missing_file = tmp_path / "missing.pdf"

    with pytest.raises(CVParsingError) as exc_info:
        extract_text_from_cv(missing_file)

    assert "CV file does not exist" in str(exc_info.value)


def test_extract_text_from_docx_preserves_table_content_order(
    tmp_path: Path,
):
    document = Document()

    document.add_paragraph("SOPHIE DUBOIS")
    document.add_paragraph(
        "Data & Business Intelligence Analyst",
    )

    document.add_paragraph("Professional Summary")
    document.add_paragraph(
        "Business intelligence analyst with six years of experience.",
    )

    document.add_paragraph("Technical Skills")
    skills_table = document.add_table(
        rows=1,
        cols=2,
    )
    skills_table.cell(0, 0).text = "Power BI\nSQL"
    skills_table.cell(0, 1).text = "Python\nPostgreSQL"

    document.add_paragraph("Languages")
    languages_table = document.add_table(
        rows=1,
        cols=2,
    )
    languages_table.cell(0, 0).text = "French: Native"
    languages_table.cell(0, 1).text = "English: Professional"

    document.add_paragraph("Certifications")
    certifications_table = document.add_table(
        rows=1,
        cols=1,
    )
    certifications_table.cell(0, 0).text = "Microsoft PL-300"

    cv_path = tmp_path / "cv-table.docx"
    document.save(cv_path)

    raw_text = extract_text_from_docx(cv_path)
    parsed_cv = parse_cv_text(raw_text)


    assert parsed_cv.full_name == "SOPHIE DUBOIS"
    assert (
        parsed_cv.professional_title
        == "Data & Business Intelligence Analyst"
    )
    assert "Power BI" in parsed_cv.skills
    assert "SQL" in parsed_cv.skills
    assert "Python" in parsed_cv.skills
    assert "PostgreSQL" in parsed_cv.skills
    assert "French: Native" in parsed_cv.languages
    assert "English: Professional" in parsed_cv.languages
    assert "Microsoft PL-300" in parsed_cv.certifications

    assert raw_text.index("Technical Skills") < raw_text.index("Power BI")
    assert raw_text.index("Power BI") < raw_text.index("Languages")


def test_parse_cv_text_maps_french_profil_to_summary():
    raw_text = """
JEAN DUPONT
Chef de projet informatique

PROFIL
Chef de projet avec 10 ans d'expérience.

COMPÉTENCES
Python, API REST
"""

    parsed_cv = parse_cv_text(raw_text)

    assert parsed_cv.summary == (
        "Chef de projet avec 10 ans d'expérience."
    )

def test_parse_cv_text_rejects_section_heading_as_full_name():
    raw_text = """
OUTILS &
LANGAGES DE PROGRAMMATION
Python
"""

    parsed_cv = parse_cv_text(raw_text)

    assert parsed_cv.full_name is None


def test_parse_cv_text_stops_skills_at_soft_skills_heading():
    raw_text = """
Alex Martin
Software Engineer

Technical Skills
Python, FastAPI, PostgreSQL

Soft Skills
Curious
Reliable
Team spirit

Languages
French: Native
"""

    parsed_cv = parse_cv_text(raw_text)

    assert parsed_cv.skills == [
        "Python",
        "FastAPI",
        "PostgreSQL",
    ]
    assert "Curious" not in parsed_cv.skills
    assert "Reliable" not in parsed_cv.skills
    assert "Team spirit" not in parsed_cv.skills


def test_parse_cv_text_merges_known_split_skill_lines():
    raw_text = """
Alex Martin
Software Engineer

Technical Skills
Cross
functional collaboration
Low -code
development

Languages
English: Professional
"""

    parsed_cv = parse_cv_text(raw_text)

    assert "Cross-functional collaboration" in parsed_cv.skills
    assert "Low-code development" in parsed_cv.skills
    assert "Cross" not in parsed_cv.skills
    assert "development" not in parsed_cv.skills


def test_parse_cv_text_does_not_treat_skill_acronyms_as_headings():
    raw_text = """
Alex Martin
Software Engineer

Technical Skills
SQL
VBA
UAT
CSS
HTML

Languages
English: Professional
"""

    parsed_cv = parse_cv_text(
        raw_text,
    )

    assert parsed_cv.skills == [
        "SQL",
        "VBA",
        "UAT",
        "CSS",
        "HTML",
    ]

    assert parsed_cv.languages == [
        "English: Professional",
    ]

    assert parsed_cv.summary is None
    assert parsed_cv.certifications == []
    assert parsed_cv.experiences == []
    
    

    
def test_extract_section_lines_collects_multiple_matching_sections():
    lines = normalize_text_lines(
        """
Alex Martin
Software Engineer

COMPETENCIES
API design
Database modelling

PROFESSIONAL EXPERIENCE
Software Engineer at Northwind Labs

TOOLS &
PROGRAMMING LANGUAGES
Python
FastAPI
PostgreSQL

LANGUAGES
French: Native
"""
    )

    skill_lines = extract_section_lines(
        lines,
        section_names=[
            "skills",
            "technical skills",
            "competencies",
            "programming languages",
            "tools & programming languages",
        ],
    )

    assert skill_lines == [
        "API design",
        "Database modelling",
        "Python",
        "FastAPI",
        "PostgreSQL",
    ]

@pytest.mark.parametrize(
    "fixture_name",
    [
        "01_pdf_single_column.pdf",
        "02_pdf_two_columns_left_sidebar.pdf",
        "03_pdf_two_columns_right_sidebar.pdf",
        "04_pdf_partial_columns.pdf",
        "05_pdf_multipage_mixed_layout.pdf",
    ],
)
def test_layout_aware_pdf_extraction_returns_usable_text(
    fixture_name: str,
):
    pdf_path = CV_PARSER_FIXTURES / fixture_name

    extracted_text = extract_text_from_pdf_pdfplumber(
        pdf_path,
    )

    assert extracted_text
    assert len(extracted_text) >= 20

def test_order_pdf_columns_places_largest_column_first():
    left_sidebar_words = [
        {
            "text": "Contact",
            "x0": 20,
            "top": 10,
        },
    ]

    right_main_words = [
        {
            "text": "Alex",
            "x0": 200,
            "top": 10,
        },
        {
            "text": "Martin",
            "x0": 240,
            "top": 10,
        },
        {
            "text": "Engineer",
            "x0": 200,
            "top": 30,
        },
    ]

    ordered_columns = order_pdf_columns(
        left_sidebar_words,
        right_main_words,
    )

    assert ordered_columns == [
        right_main_words,
        left_sidebar_words,
    ]



def test_layout_aware_extraction_preserves_left_sidebar_sections():
    pdf_path = (
        CV_PARSER_FIXTURES
        / "02_pdf_two_columns_left_sidebar.pdf"
    )

    extracted_text = extract_text_from_pdf_pdfplumber(
        pdf_path,
    )
    parsed_cv = parse_cv_text(
        extracted_text,
    )

    assert parsed_cv.full_name == "Alex Martin"
    assert parsed_cv.professional_title == "Software Engineer"

    assert "Python" in parsed_cv.skills
    assert "FastAPI" in parsed_cv.skills
    assert "PostgreSQL" in parsed_cv.skills
    assert "React" in parsed_cv.skills
    assert "Docker" in parsed_cv.skills
    assert "SQL" in parsed_cv.skills

    assert "French: Native" in parsed_cv.languages
    assert "English: Professional" in parsed_cv.languages

    assert parsed_cv.certifications == [
        "Microsoft Azure Fundamentals",
    ]

    assert "Alex Martin" not in parsed_cv.certifications
    assert "Software Engineer" not in parsed_cv.certifications

def test_layout_aware_extraction_preserves_right_sidebar_sections():
    pdf_path = (
        CV_PARSER_FIXTURES
        / "03_pdf_two_columns_right_sidebar.pdf"
    )

    extracted_text = extract_text_from_pdf_pdfplumber(
        pdf_path,
    )
    parsed_cv = parse_cv_text(
        extracted_text,
    )

    assert parsed_cv.full_name == "Alex Martin"
    assert parsed_cv.professional_title == "Software Engineer"

    assert "Python" in parsed_cv.skills
    assert "FastAPI" in parsed_cv.skills
    assert "PostgreSQL" in parsed_cv.skills
    assert "React" in parsed_cv.skills
    assert "Docker" in parsed_cv.skills
    assert "SQL" in parsed_cv.skills

    assert "French: Native" in parsed_cv.languages
    assert "English: Professional" in parsed_cv.languages

    assert "Microsoft Azure Fundamentals" in parsed_cv.certifications

    assert "Python" not in parsed_cv.languages
    assert "React" not in parsed_cv.languages


def test_pdf_extraction_quality_penalizes_language_overflow():
    valid_text = """
Alex Martin
Software Engineer

Technical Skills
Python
FastAPI
PostgreSQL

Languages
French: Native
English: Professional

Professional Experience
Software Engineer at Northwind Labs
Built reliable APIs.
"""

    interleaved_text = (
        CV_PARSER_FIXTURES
        / "09_pdf_interleaved_internal_order.txt"
    ).read_text(
        encoding="utf-8",
    )

    valid_quality = calculate_pdf_extraction_quality(
        valid_text,
    )
    interleaved_quality = calculate_pdf_extraction_quality(
        interleaved_text,
    )

    assert valid_quality > interleaved_quality


def test_select_best_pdf_extraction_prefers_structured_result():
    structured_text = """
Alex Martin
Software Engineer

Technical Skills
Python
FastAPI
PostgreSQL

Languages
French: Native
English: Professional

Professional Experience
Software Engineer at Northwind Labs
Built reliable APIs.
"""

    interleaved_text = (
        CV_PARSER_FIXTURES
        / "09_pdf_interleaved_internal_order.txt"
    ).read_text(
        encoding="utf-8",
    )

    selected_text = select_best_pdf_extraction(
        [
            interleaved_text,
            structured_text,
        ]
    )

    assert selected_text == structured_text.strip()


def test_pdf_extraction_orchestrator_parses_sidebar_fixture():
    pdf_path = (
        CV_PARSER_FIXTURES
        / "03_pdf_two_columns_right_sidebar.pdf"
    )

    extracted_text = extract_text_from_pdf(
        pdf_path,
    )
    parsed_cv = parse_cv_text(
        extracted_text,
    )

    assert parsed_cv.full_name == "Alex Martin"
    assert parsed_cv.skills
    assert parsed_cv.languages
    assert parsed_cv.experiences


def test_pypdf2_extractor_remains_available_as_fallback():
    pdf_path = (
        CV_PARSER_FIXTURES
        / "01_pdf_single_column.pdf"
    )

    extracted_text = extract_text_from_pdf_pypdf2(
        pdf_path,
    )

    assert extracted_text
    assert "Alex Martin" in extracted_text
    
    
def test_merge_split_french_section_headings():
    lines = [
        "OUTILS &",
        "LANGAGES DE",
        "PROGRAMMATION",
        "Python",
        "QUALITÉS &",
        "LANGUES",
        "Français : bilingue",
    ]

    merged_lines = merge_split_section_headings(
        lines,
    )

    assert merged_lines == [
        "Outils & Langages de Programmation",
        "Python",
        "Qualités & Langues",
        "Français : bilingue",
    ]


def test_extract_languages_filters_mixed_qualities_section():
    raw_text = """
QUALITÉS &
LANGUES
Curieux
Organisé
Français : bilingue
Analytique
Anglais : courant
Créatif
Tamoul : bilingue
Minutieux
Collaboratif
"""

    lines = merge_split_section_headings(
        normalize_text_lines(raw_text),
    )

    languages = extract_languages(
        lines,
    )

    assert languages == [
        "Français : bilingue",
        "Anglais : courant",
        "Tamoul : bilingue",
    ]


def test_parse_cv_text_supports_french_profile_sections():
    raw_text = """
PROFIL
Développeur web à la recherche d'une alternance.

COMPETENCES
Conception de bases de données
Déploiement de sites web

EXPERIENCES PROFESSIONNELLES
DÉVELOPPEUR WEB 2023
Création et déploiement d'un site web.

LATHAN
T A R M A T
Étudiant en informatique
à l'école EPITECH

OUTILS &
LANGAGES DE
PROGRAMMATION
Python
PHP
C
C++
HTML
Java
CSS
SQL

QUALITÉS &
LANGUES
Curieux
Organisé
Français : bilingue
Anglais : courant
Tamoul : bilingue
"""

    parsed_cv = parse_cv_text(
        raw_text,
    )

    assert parsed_cv.full_name == "Lathan Tarmat"
    assert (
        parsed_cv.professional_title
        == "Étudiant en informatique"
    )

    assert "Conception de bases de données" in parsed_cv.skills
    assert "Déploiement de sites web" in parsed_cv.skills
    assert "Python" in parsed_cv.skills
    assert "PHP" in parsed_cv.skills
    assert "HTML" in parsed_cv.skills
    assert "SQL" in parsed_cv.skills

    assert parsed_cv.languages == [
        "Français : bilingue",
        "Anglais : courant",
        "Tamoul : bilingue",
    ]

    assert "Curieux" not in parsed_cv.languages
    assert "Organisé" not in parsed_cv.languages

    assert parsed_cv.experiences
    

def test_parse_cv_text_reconstructs_multiline_experiences():
    raw_text = """
Alex Martin
Software Engineer

PROFESSIONAL EXPERIENCE
SOFTWARE ENGINEER JANUARY 2022
NORTHWIND LABS DECEMBER 2024
Built APIs and automated deployment pipelines.

WEB DEVELOPER 2020
CONTOSO STUDIO 2021
Developed web interfaces.
"""

    parsed_cv = parse_cv_text(
        raw_text,
    )

    assert len(parsed_cv.experiences) == 2

    first_experience = parsed_cv.experiences[0]

    assert first_experience.title == "SOFTWARE ENGINEER"
    assert first_experience.company == "NORTHWIND LABS"
    assert first_experience.start_date == "JANUARY 2022"
    assert first_experience.end_date == "DECEMBER 2024"
    assert first_experience.description == (
        "Built APIs and automated deployment pipelines."
    )

    second_experience = parsed_cv.experiences[1]

    assert second_experience.title == "WEB DEVELOPER"
    assert second_experience.company == "CONTOSO STUDIO"
    assert second_experience.start_date == "2020"
    assert second_experience.end_date == "2021"
    assert second_experience.description == (
        "Developed web interfaces."
    )
    
def test_parse_cv_text_reconstructs_french_project_experiences():
    raw_text = """
Lathan Tarmat
Étudiant en informatique

EXPERIENCES PROFESSIONNELLES
DÉVELOPPEUR WEB MAI 2023
LE PRO DE L'IMMO SEPTEMBRE 2023
Étude des besoins.
Développement du site.
Mise en ligne.

PROJET INFORMATIQUE DE L3 26/01/2023
APPLICATION DE GESTION DE STOCK 18/04/2023
Université Paris Cité
Application permettant de gérer les stocks.

PROJET INFORMATIQUE DE L2 17/01/2022
CAPTCHA 18/04/2022
Université Paris Cité
Réalisation d'un système d'authentification.

PROFESSEUR PARTICULIER SEPTEMBRE 2017
PARIS, IVRY-SUR-SEINE JANVIER 2023
Adaptation des cours au niveau de l'élève.

FORMATION
Licence Informatique
"""

    parsed_cv = parse_cv_text(
        raw_text,
    )

    assert len(parsed_cv.experiences) == 4

    web_experience = parsed_cv.experiences[0]

    assert web_experience.title == "DÉVELOPPEUR WEB"
    assert web_experience.company == "LE PRO DE L'IMMO"
    assert web_experience.start_date == "MAI 2023"
    assert web_experience.end_date == "SEPTEMBRE 2023"
    assert web_experience.description == (
        "Étude des besoins. "
        "Développement du site. "
        "Mise en ligne."
    )

    stock_project = parsed_cv.experiences[1]

    assert stock_project.title == (
        "APPLICATION DE GESTION DE STOCK"
    )
    assert stock_project.company == "Université Paris Cité"
    assert stock_project.start_date == "26/01/2023"
    assert stock_project.end_date == "18/04/2023"
    assert stock_project.description == (
        "Application permettant de gérer les stocks."
    )

    captcha_project = parsed_cv.experiences[2]

    assert captcha_project.title == "CAPTCHA"
    assert captcha_project.company == "Université Paris Cité"
    assert captcha_project.description == (
        "Réalisation d'un système d'authentification."
    )

    teaching_experience = parsed_cv.experiences[3]

    assert teaching_experience.title == (
        "PROFESSEUR PARTICULIER"
    )
    assert teaching_experience.company is None
    assert teaching_experience.start_date == (
        "SEPTEMBRE 2017"
    )
    assert teaching_experience.end_date == (
        "JANVIER 2023"
    )
    assert "PARIS, IVRY-SUR-SEINE" in (
        teaching_experience.description or ""
    )